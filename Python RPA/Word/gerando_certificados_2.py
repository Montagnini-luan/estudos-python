from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from openpyxl import load_workbook

nome_arquivo_alunos = "C:\\Users\\luanm\\OneDrive\\Desktop\\Sites\\Estudos\\estudos python\\Python RPA\\Word\\DadosAlunos.xlsx"

planilha_dados_alunos = load_workbook(nome_arquivo_alunos)

sheet_selecionada = planilha_dados_alunos['Nomes']

for linha in range(2, len(sheet_selecionada["A"]) + 1):

    arquivo_word = Document("C:\\Users\\luanm\\OneDrive\\Desktop\\Sites\\Estudos\\estudos python\\Python RPA\\Word\\Certificado2.docx")

    estilo = arquivo_word.styles["Normal"]

    nome_aluno = sheet_selecionada['A%s' % linha].value
    dia = sheet_selecionada['B%s' % linha].value
    mes = sheet_selecionada['C%s' % linha].value
    ano = sheet_selecionada['D%s' % linha].value
    curso = sheet_selecionada['E%s' % linha].value
    instrutor = sheet_selecionada['F%s' % linha].value

    frase_parte1 = "Concluiu com sucesso o curso de "
    frase_parte2 = ", com a carga horária de 20 horas, promovido pela escola de Cursos Online em "
    frase_montada = f"{frase_parte2} {dia} de {mes} de {ano}"


    for paragrafo in arquivo_word.paragraphs:

        if "@nome" in paragrafo.text:

            paragrafo.text = nome_aluno
            fonte = estilo.font
            fonte.name = "Calibri (Corpo)"
            fonte.size = Pt(24)

        if "Dezembro" in paragrafo.text:

            paragrafo.text = frase_parte1
            fonte = estilo.font
            fonte.name = "Calibri (Corpo)"
            fonte.size = Pt(24)
            adicona_palavra = paragrafo.add_run(curso)
            adicona_palavra.font.color.rgb = RGBColor(255, 0, 0)
            adicona_palavra.underline = True
            adicona_palavra.bold = True
            adicona_palavra = paragrafo.add_run(frase_montada)
            adicona_palavra.font.color.rgb = RGBColor(0, 0, 0)

        if "Instrutor" in paragrafo.text:

            paragrafo.text = instrutor + " - Instrutor"
            fonte = estilo.font
            fonte.name = "Calibri (Corpo)"
            fonte.size = Pt(24)

    caminho_certificados = "C:\\Users\\luanm\OneDrive\\Desktop\\Sites\\Estudos\\estudos python\\Python RPA\\Word\\Certificados\\" + nome_aluno + ".docx"

    arquivo_word.save(caminho_certificados)

print("Certificados gerados com sucesso")
